"""
export_word.py - blueprint-style Word (.docx) export helpers
Generates .docx for single learner and streams to client using BytesIO.
Safe: no leftover temp files, works on Windows/Linux.
"""
import os
import logging
from io import BytesIO
from flask import Blueprint, send_file, abort
from docx import Document
from docx.shared import Pt

from utils import get_learner_by_id

logger = logging.getLogger(__name__)

export_word_bp = Blueprint("export_word", __name__, url_prefix="/export")



def _build_document_for_learner(learner: dict) -> Document:
    """
    Convert learner dict into a python-docx Document object.
    Adjust the fields used below to match your learner dictionary returned by utils.
    """
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(11)

    # Header / basic info
    doc.add_heading(learner.get("full_name", "Unnamed Learner"), level=1)
    doc.add_paragraph(f"Admission No: {learner.get('admission_no', 'N/A')}")
    doc.add_paragraph(f"Grade: {learner.get('grade', 'N/A')}")
    doc.add_paragraph(f"Term: {learner.get('term', 'N/A')}")
    doc.add_paragraph(" ")

    # Example: summary table for marks (if present)
    marks = learner.get("marks", [])
    if marks:
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Subject"
        hdr_cells[1].text = "Mark"
        hdr_cells[2].text = "Grade"
        for m in marks:
            row_cells = table.add_row().cells
            row_cells[0].text = m.get("subject", "")
            row_cells[1].text = str(m.get("score", ""))
            row_cells[2].text = m.get("grade", "")
    else:
        doc.add_paragraph("No marks available.")

    # Footer / comments
    doc.add_paragraph("Comments:")
    doc.add_paragraph(learner.get("comments", ""))

    return doc


@export_word_bp.route("/learner/<int:learner_id>")
def export_learner_word(learner_id: int):
    """Route: /export/learner/<id>.docx  - returns a .docx file"""
    learner = get_learner_by_id(learner_id)
    if not learner:
        abort(404, description="Learner not found")

    doc = _build_document_for_learner(learner)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    filename = f"{learner.get('full_name', 'learner')}_report.docx"
    return send_file(
        bio,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
